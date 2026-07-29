from __future__ import annotations

import argparse
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader, PdfWriter


AUTHOR = "Neil Mitchell"
ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DOCX = ROOT / "documents" / "Neil-Mitchell-Resume.docx"
DEFAULT_PDF = ROOT / "output" / "pdf" / "Neil-Mitchell-Resume.pdf"


def set_font(font, name: str, size: int, *, bold: bool | None = None) -> None:
    font.name = name
    font.size = Pt(size)
    font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        font.bold = bold


def set_run_font(run, name: str = "Calibri", size: int = 11) -> None:
    set_font(run.font, name, size)
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_style_font(style, name: str = "Calibri", size: int = 11) -> None:
    set_font(style.font, name, size)
    style.element.get_or_add_rPr()
    style.element.rPr.rFonts.set(qn("w:ascii"), name)
    style.element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), name)


def configure_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    set_style_font(normal)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.04

    heading_one = styles["Heading 1"]
    set_style_font(heading_one)
    heading_one.font.bold = True
    heading_one.font.all_caps = True
    heading_one.paragraph_format.space_before = Pt(9)
    heading_one.paragraph_format.space_after = Pt(6)
    heading_one.paragraph_format.keep_with_next = True
    heading_one.paragraph_format.keep_together = True

    heading_two = styles["Heading 2"]
    set_style_font(heading_two)
    heading_two.font.bold = True
    heading_two.paragraph_format.space_before = Pt(6)
    heading_two.paragraph_format.space_after = Pt(4)
    heading_two.paragraph_format.keep_with_next = True
    heading_two.paragraph_format.keep_together = True

    bullet = styles["List Bullet"]
    set_style_font(bullet)
    bullet.paragraph_format.left_indent = Inches(0.25)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_before = Pt(0)
    bullet.paragraph_format.space_after = Pt(0.75)
    bullet.paragraph_format.line_spacing = 1.04

    if "Resume Metadata" not in styles:
        metadata = styles.add_style("Resume Metadata", WD_STYLE_TYPE.PARAGRAPH)
    else:
        metadata = styles["Resume Metadata"]
    set_style_font(metadata, size=10)
    metadata.paragraph_format.space_after = Pt(0)
    metadata.paragraph_format.line_spacing = 1


def add_paragraph(
    document: Document,
    text: str,
    *,
    size: int = 11,
    bold: bool = False,
    italic: bool = False,
    space_before: float = 0,
    space_after: float = 0,
    keep_with_next: bool = False,
    style: str | None = None,
):
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.line_spacing = 1.04
    run = paragraph.add_run(text)
    set_run_font(run, size=size)
    run.bold = bold
    run.italic = italic
    return paragraph


def add_hyperlink(
    paragraph,
    text: str,
    url: str,
    *,
    name: str = "Calibri",
    size: int = 10,
) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    run_fonts = OxmlElement("w:rFonts")
    run_fonts.set(qn("w:ascii"), name)
    run_fonts.set(qn("w:hAnsi"), name)
    run_fonts.set(qn("w:eastAsia"), name)
    run_properties.append(run_fonts)

    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), str(size * 2))
    run_properties.append(font_size)

    font_size_complex = OxmlElement("w:szCs")
    font_size_complex.set(qn("w:val"), str(size * 2))
    run_properties.append(font_size_complex)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    run_properties.append(underline)

    run.append(run_properties)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_role(
    document: Document,
    title: str,
    organisation_and_dates: str,
    context: str,
    bullets: list[str],
) -> None:
    title_paragraph = document.add_paragraph(title, style="Heading 2")
    for run in title_paragraph.runs:
        set_run_font(run)
        run.bold = True

    add_paragraph(
        document,
        organisation_and_dates,
        italic=True,
        space_after=5,
        keep_with_next=True,
    )
    add_paragraph(document, context, italic=True, space_after=5, keep_with_next=True)

    for bullet_text in bullets:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.keep_together = True
        run = paragraph.add_run(bullet_text)
        set_run_font(run)


def build_resume(output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    configure_styles(document)

    properties = document.core_properties
    properties.author = AUTHOR
    properties.last_modified_by = AUTHOR
    properties.title = "Neil Mitchell Resume"
    properties.subject = "Professional resume"
    properties.keywords = (
        "project management, applied AI, machine learning, LLM workflows, automation, "
        "software development, insurance technology"
    )
    properties.comments = "Updated July 2026"
    properties.created = datetime(2026, 7, 26, tzinfo=timezone.utc)
    properties.modified = datetime(2026, 7, 29, tzinfo=timezone.utc)
    properties.revision = 4

    add_paragraph(document, AUTHOR, size=16, bold=True, space_after=8, keep_with_next=True)
    add_paragraph(
        document,
        "Project Manager II - Applied AI/ML Engineering | TD Bank Group",
        bold=True,
        space_after=8,
        keep_with_next=True,
    )
    contact = document.add_paragraph(style="Resume Metadata")
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(16)
    contact.paragraph_format.keep_with_next = True
    contact.paragraph_format.line_spacing = 1
    contact_details = contact.add_run(
        "Saint John, New Brunswick | 506-639-9083 | "
        "neil_mitchell89@hotmail.com | "
    )
    set_run_font(contact_details, size=10)
    add_hyperlink(contact, "neilmitchell.ca", "https://neilmitchell.ca")

    document.add_paragraph("PROFESSIONAL SUMMARY", style="Heading 1")
    add_paragraph(
        document,
        "Project and delivery leader coordinating applied AI/ML engineering work that supports "
        "TD Insurance. Experienced translating complex business and technical objectives into "
        "structured plans, clear ownership, and visible priorities, risks, dependencies, and "
        "milestones. Brings seven years of insurance context across claims, vendor operations, "
        "digital servicing, product and regulatory work, and Guidewire-based platforms.",
        space_after=11,
    )

    document.add_paragraph("WORK EXPERIENCE", style="Heading 1")

    add_role(
        document,
        "Project Manager II - Applied AI/ML Engineering",
        "AI2, TD Bank Group | 07/2026 - Present",
        "Supporting TD Insurance",
        [
            "Lead planning and execution for applied AI/ML engineering work supporting TD Insurance.",
            "Coordinate priorities, dependencies, risks, decisions, and milestones across "
            "engineering and business workstreams.",
            "Manage work through a Kanban model, maintaining clear ownership and visibility across "
            "active priorities.",
            "Partner with engineers, senior leaders, product partners, data specialists, and "
            "business teams to turn objectives into actionable delivery plans.",
            "Support implementation, integration, testing, validation, production readiness, "
            "monitoring, and operational handoff.",
        ],
    )

    add_role(
        document,
        "Senior Product Analyst",
        "GIJ, TD Insurance | 08/2025 - 07/2026",
        "Strategic Product & Regulatory Knowledge (SPARK)",
        [
            "Shaped high-level product requirements and feature definitions for Guidewire "
            "PolicyCenter initiatives, including AI RSP and DASH.",
            "Coordinated scope, sequencing, acceptance criteria, and ownership across multiple "
            "workstreams.",
            "Partnered with business analysts and platform teams to prepare implementation-ready "
            "stories.",
            "Managed product, technology, and data trade-offs while contributing to roadmap "
            "planning and delivery readiness.",
        ],
    )

    page_break = document.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)

    add_role(
        document,
        "Senior Product Analyst",
        "GIJ, TD Insurance | 02/2024 - 08/2025",
        "Digital Service Performance (MyInsurance)",
        [
            "Owned high-level product requirements and feature definitions for customer servicing "
            "journeys within MyInsurance.",
            "Led backlog refinement and story shaping with business analysts.",
            "Coordinated work across teams and prepared MyInsurance enhancements for release.",
            "Managed sequencing and risk with technology and operations partners while building "
            "deep knowledge of digital insurance servicing.",
        ],
    )

    add_role(
        document,
        "Vendor Analyst II",
        "CFLVS, TD Insurance | 04/2020 - 02/2024",
        "Vendor Management Office (VMO)",
        [
            "Managed an enterprise Accounts Receivable process on a quarterly cadence, processing "
            "80,000+ invoices totaling $55MM+ in outstanding payments.",
            "Designed and implemented VBA automation pipelines for data aggregation and "
            "reconciliation, improving cycle time and accuracy.",
            "Led weekly progress and delivery reviews with stakeholders across multiple business "
            "lines.",
            "Recognized as a TDI All Star recipient in 2022 for delivery excellence and stakeholder "
            "leadership.",
        ],
    )

    add_role(
        document,
        "Claims Advisor",
        "CFLVS, TD Insurance | 04/2019 - 04/2020",
        "Claims Response Office (CRO)",
        [
            "Handled inbound FNOL and outbound accident benefits calls, opening exposures and "
            "preparing claims for initial adjuster contact.",
            "Built foundational knowledge of claims intake workflows, accident benefits processes, "
            "and operational handoff to adjusting teams.",
        ],
    )

    document.add_paragraph("CORE PRODUCT & DELIVERY SKILLS", style="Heading 1")
    add_paragraph(
        document,
        "Applied AI/ML Delivery | Project Management | Kanban | Product Ownership & Backlog "
        "Shaping | Cross-Team Delivery Leadership | Scope & Risk Management | Stakeholder & "
        "Executive Communication | Roadmap & Milestone Planning | Requirements Engineering & "
        "Acceptance Criteria",
        space_after=6,
    )

    document.add_paragraph("TECHNICAL SKILLS", style="Heading 1")
    add_paragraph(
        document,
        "AI & Automation: LLM and agent workflows, OpenAI and Anthropic API integration, "
        "prompt and tool design, AI-assisted development (Codex and GitHub Copilot), Python "
        "and VBA automation",
        size=10,
    )
    add_paragraph(
        document,
        "Development: Python, TypeScript/JavaScript, Kotlin, Java, C#, C++, VBA, Lua; "
        "Next.js, React, FastAPI, Jetpack Compose, GitHub Actions, Docker, Vercel",
        size=10,
    )
    add_paragraph(
        document,
        "Enterprise & Data: Guidewire PolicyCenter, ClaimCenter, BillingCenter; Jira, "
        "Confluence, GitHub; SQL, PostgreSQL, MySQL, Microsoft SQL Server, Azure SQL, "
        "SQLite, Excel, Access, Tableau",
        size=10,
        space_after=6,
    )

    document.add_paragraph("INTERESTS", style="Heading 1")
    add_paragraph(
        document,
        "Sports | Fitness | Travel | Family | Product Strategy | Software Development | "
        "Community Volunteering | Mentorship | Disability & Inclusion",
        size=10,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def finalize_pdf(input_path: Path, output_path: Path) -> None:
    reader = PdfReader(input_path)
    writer = PdfWriter()
    writer.clone_document_from_reader(reader)

    metadata = {
        key: str(value)
        for key, value in (reader.metadata or {}).items()
        if value is not None and key not in {"/Author", "/Creator", "/Producer"}
    }
    metadata.update(
        {
            "/Title": "Neil Mitchell Resume",
            "/Subject": "Professional resume",
            "/Author": AUTHOR,
            "/Creator": AUTHOR,
            "/Producer": AUTHOR,
            "/LastModifiedBy": AUTHOR,
            "/Modifier": AUTHOR,
            "/ModDate": "D:20260729000000-03'00'",
        }
    )
    writer.add_metadata(metadata)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("wb") as output_stream:
        writer.write(output_stream)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build and finalize Neil Mitchell's resume.")
    parser.add_argument("--docx-output", type=Path, default=DEFAULT_DOCX)
    parser.add_argument("--pdf-input", type=Path)
    parser.add_argument("--pdf-output", type=Path, default=DEFAULT_PDF)
    args = parser.parse_args()

    build_resume(args.docx_output.resolve())
    if args.pdf_input:
        finalize_pdf(args.pdf_input.resolve(), args.pdf_output.resolve())


if __name__ == "__main__":
    main()
